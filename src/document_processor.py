"""
DocMemory - Document Processing Pipeline
Handles various document formats and content extraction
"""
import os
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Tuple
from dataclasses import dataclass
import hashlib

try:
    import PyPDF2
    from pdfminer.high_level import extract_text as pdf_extract_text
except ImportError:
    PyPDF2 = None
    pdf_extract_text = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

@dataclass
class DocumentChunk:
    """Represents a chunk of document content for processing"""
    content: str
    page_number: int = 1
    chunk_index: int = 0
    metadata: Dict[str, Any] = None

class DocumentProcessor:
    """Processes various document formats and extracts content"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.csv': self._process_csv,
            '.html': self._process_html,
            '.rtf': self._process_txt,  # Treat RTF as text for simplicity
            '.odt': self._process_txt,  # Treat ODT as text for simplicity
        }
        
        # Maximum chunk size in characters
        self.max_chunk_size = 1000
        self.chunk_overlap = 100
        self.default_title = "Untitled Document"
    
    def process_document(self, file_path: str, title: str = None) -> List[DocumentChunk]:
        """Main method to process a document based on its format"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")
        
        # Get document title
        if title is None:
            title = file_path.stem
        
        # Process document based on format
        process_func = self.supported_formats[extension]
        content = process_func(file_path)
        
        # Create chunks from content
        chunks = self._create_chunks(content, title)
        
        # Add metadata to each chunk
        for i, chunk in enumerate(chunks):
            chunk.metadata = {
                'source_file': str(file_path),
                'document_type': extension[1:],  # Remove the dot
                'chunk_count': len(chunks),
                'total_size': len(content),
                'title': title
            }
        
        return chunks
    
    def _process_pdf(self, file_path: Path) -> str:
        """Process PDF files"""
        if pdf_extract_text:
            try:
                # Try pdfminer first (better text extraction)
                text = pdf_extract_text(str(file_path))
                return text.strip()
            except Exception:
                pass
        
        # Fallback to PyPDF2
        if PyPDF2:
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    return text.strip()
            except Exception as e:
                print(f"Error processing PDF with PyPDF2: {e}")
        
        # If neither library works, raise an error
        raise Exception("Failed to process PDF document. Install PyPDF2 or pdfminer.six.")
    
    def _process_docx(self, file_path: Path) -> str:
        """Process DOCX files"""
        if Document is None:
            raise Exception("docx library not available. Install python-docx to process DOCX files.")
        
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            return '\n'.join(text).strip()
        except Exception as e:
            raise Exception(f"Error processing DOCX file: {e}")
    
    def _process_txt(self, file_path: Path) -> str:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read().strip()
        except Exception as e:
            raise Exception(f"Error processing text file: {e}")
    
    def _process_csv(self, file_path: Path) -> str:
        """Process CSV files"""
        if pd is None:
            raise Exception("pandas library not available. Install pandas to process CSV files.")
        
        try:
            df = pd.read_csv(file_path)
            # Convert to string representation
            text = df.to_string(index=False)
            return text.strip()
        except Exception as e:
            raise Exception(f"Error processing CSV file: {e}")
    
    def _process_html(self, file_path: Path) -> str:
        """Process HTML files"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise Exception("BeautifulSoup4 library not available. Install beautifulsoup4 to process HTML files.")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text()
                # Clean up text
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                return text.strip()
        except Exception as e:
            raise Exception(f"Error processing HTML file: {e}")
    
    def _create_chunks(self, content: str, title: str) -> List[DocumentChunk]:
        """Split content into manageable chunks"""
        chunks = []
        
        # Split content into chunks of max_chunk_size with overlap
        start = 0
        chunk_index = 0
        
        while start < len(content):
            end = start + self.max_chunk_size
            
            # If we're not at the end, try to break at sentence boundary
            if end < len(content):
                # Look for a good breaking point (sentence end, paragraph end, or space)
                search_start = end - self.chunk_overlap
                break_point = end
                
                # Look for punctuation followed by space
                for i in range(min(end, len(content)) - 1, search_start, -1):
                    if content[i] in '.!?;':
                        if i + 1 < len(content) and content[i + 1] in ' \n\t':
                            break_point = i + 1
                            break
                        elif i + 2 < len(content) and content[i + 1] in ' \n\t':
                            break_point = i + 2
                            break
                
                # If no good breaking point found, use the overlap point
                if break_point == end:
                    break_point = max(search_start, start + 50)  # Ensure minimum chunk size
            
            chunk_content = content[start:break_point].strip()
            if chunk_content:  # Only add non-empty chunks
                chunks.append(DocumentChunk(
                    content=chunk_content,
                    page_number=1,  # Will be updated if processing multi-page docs
                    chunk_index=chunk_index
                ))
                chunk_index += 1
            
            start = break_point
        
        return chunks

class DocumentIngestionPipeline:
    """Main pipeline for ingesting documents into DocMemory"""
    
    def __init__(self, docmemory_system):
        self.docmemory_system = docmemory_system
        self.processor = DocumentProcessor()
        
        # Embedding model placeholder (will be set externally)
        self.embedding_model = None
    
    def set_embedding_model(self, model):
        """Set the embedding model for processing"""
        self.embedding_model = model
    
    def process_and_store_document(self, 
                                   file_path: str, 
                                   title: str = None,
                                   tags: List[str] = None,
                                   custom_metadata: Dict[str, Any] = None) -> List[str]:
        """Process a document and store it in DocMemory"""
        if self.embedding_model is None:
            raise ValueError("Embedding model must be set before processing documents")
        
        # Process the document
        chunks = self.processor.process_document(file_path, title)
        
        # Store each chunk as a separate memory
        stored_ids = []
        
        for chunk in chunks:
            # Generate embedding for the chunk content
            embedding = self.embedding_model.encode([chunk.content])[0]
            
            # Create metadata combining document metadata and chunk info
            metadata = {**chunk.metadata}
            if custom_metadata:
                metadata.update(custom_metadata)
            
            # Store in DocMemory
            doc_id = self.docmemory_system.add_document(
                content=chunk.content,
                title=chunk.metadata.get('title', self.processor.default_title),
                source_file=chunk.metadata['source_file'],
                embedding=embedding,
                document_type=chunk.metadata['document_type'],
                tags=tags or [],
                metadata=metadata,
                summary="",  # Will be generated later if needed
                page_numbers=[chunk.page_number]
            )
            
            stored_ids.append(doc_id)
        
        print(f"Successfully processed and stored {len(stored_ids)} document chunks from {file_path}")
        return stored_ids
    
    def batch_process_documents(self, 
                                file_paths: List[str],
                                tags_by_file: Dict[str, List[str]] = None,
                                metadata_by_file: Dict[str, Dict[str, Any]] = None) -> Dict[str, List[str]]:
        """Process multiple documents at once"""
        results = {}
        
        for file_path in file_paths:
            try:
                tags = tags_by_file.get(file_path) if tags_by_file else None
                metadata = metadata_by_file.get(file_path) if metadata_by_file else None
                
                stored_ids = self.process_and_store_document(
                    file_path=file_path,
                    tags=tags,
                    custom_metadata=metadata
                )
                
                results[file_path] = stored_ids
            except Exception as e:
                print(f"Error processing {file_path}: {e}")
                results[file_path] = []
        
        return results
    
    def update_document(self, 
                        file_path: str,
                        doc_ids: List[str],
                        new_title: str = None,
                        new_tags: List[str] = None) -> bool:
        """Update an existing document with new content"""
        try:
            # Process the updated document
            chunks = self.processor.process_document(file_path, new_title)
            
            # For now, we'll delete old chunks and add new ones
            # In a more sophisticated system, we might want to match and update specific chunks
            for doc_id in doc_ids:
                self.docmemory_system.core_memory.delete_document(doc_id)
            
            # Process and store the new content
            new_doc_ids = self.process_and_store_document(
                file_path=file_path,
                title=new_title,
                tags=new_tags
            )
            
            print(f"Successfully updated document: {file_path}")
            return True
        except Exception as e:
            print(f"Error updating document {file_path}: {e}")
            return False

# Example usage
if __name__ == "__main__":
    # This would be used with the DocMemory system
    print("Document processor ready.")